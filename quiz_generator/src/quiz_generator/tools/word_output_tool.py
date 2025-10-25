import json
from docx import Document
from docx.shared import Inches
# --- Corrected Import for BaseTool (Attempt 2) ---
from crewai.tools import BaseTool
# ---
from pydantic import BaseModel, Field

class WordOutputToolInput(BaseModel):
    """Input schema for WordOutputTool."""
    filename: str = Field(..., description="The desired filename for the output Word document (e.g., 'final_quiz'). .docx extension will be added if missing.")
    quiz_data_json: str = Field(..., description="A JSON string representing the final, approved quiz data, containing 'multiple_choice' and 'true_false' keys.")

class WordOutputTool(BaseTool):
    name: str = "Word Output Tool"
    description: str = "Formats quiz data (MCQ and T/F) from a JSON string into a structured Word document (.docx) and saves it to a file."
    args_schema: type[BaseModel] = WordOutputToolInput

    def _run(self, filename: str, quiz_data_json: str) -> str:
        """Formats and saves quiz data to a Word file."""
        try:
            quiz_data = json.loads(quiz_data_json)
        except json.JSONDecodeError as e:
            return f"Error: Invalid JSON string provided. Details: {e}"
        except Exception as e:
            return f"Error parsing JSON data: {e}"

        if not isinstance(quiz_data, dict):
             return "Error: Parsed quiz data is not a dictionary."

        mcq_questions = quiz_data.get("multiple_choice", [])
        tf_questions = quiz_data.get("true_false", [])

        if not isinstance(mcq_questions, list):
             mcq_questions = []
             print("Warning: 'multiple_choice' key did not contain a list.")
        if not isinstance(tf_questions, list):
             tf_questions = []
             print("Warning: 'true_false' key did not contain a list.")


        doc = Document()
        doc.add_heading('Generated Quiz', level=1)

        if mcq_questions:
            doc.add_heading('Multiple Choice Questions', level=2)
            for i, q in enumerate(mcq_questions, 1):
                if not isinstance(q, dict): continue
                question_text = q.get("question", "Missing question text")
                options = q.get("options", [])
                correct_answer = q.get("correct_answer", "Missing correct answer")

                doc.add_paragraph(f"{i}. {question_text}", style='List Number')
                if isinstance(options, list):
                    option_labels = ["A", "B", "C", "D"]
                    for j, option in enumerate(options):
                         if j < len(option_labels):
                             doc.add_paragraph(f"   {option_labels[j]}) {option}", style='List Bullet 2')
                         else:
                             doc.add_paragraph(f"   - {option}", style='List Bullet 2')
                doc.add_paragraph(f"   Correct Answer: {correct_answer}\n")

        if tf_questions:
            doc.add_heading('True/False Questions', level=2)
            for i, q in enumerate(tf_questions, 1):
                 if not isinstance(q, dict): continue
                 question_text = q.get("question", "Missing question text")
                 answer = q.get("answer", "Missing answer")

                 doc.add_paragraph(f"{i}. {question_text}", style='List Number')
                 doc.add_paragraph(f"   Answer: {str(answer).capitalize()}\n")

        if not filename.lower().endswith(".docx"):
            filename += ".docx"

        try:
            doc.save(filename)
            return f"Successfully saved formatted quiz to Word document '{filename}'"
        except Exception as e:
            return f"Error saving Word document '{filename}': {e}"

