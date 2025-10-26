# adk_quiz_generator/src/adk_quiz_generator/tools/word_tools.py

from google.adk.tools import FunctionTool
from docx import Document
import json

def quiz_to_word_with_sources(file_path: str, quiz_json_str: str) -> str:
    """
    Generates a Word document from the final quiz JSON object.
    Handles a dictionary containing MCQs, T/F questions, validation notes,
    and fact-checking sources.
    """
    try:
        quiz_data = json.loads(quiz_json_str)
    except json.JSONDecodeError as e:
        return f"Invalid JSON: {e}"

    if not isinstance(quiz_data, dict):
        return "Expected a JSON object (dict) at the top level."

    doc = Document()
    doc.add_heading("Generated Quiz", level=0)

    # Get the lists of questions
    mcq_questions = quiz_data.get("multiple_choice", [])
    tf_questions = quiz_data.get("true_false", [])
    all_questions = mcq_questions + tf_questions

    if not all_questions:
        return "No questions found in the JSON."

    # --- Process all questions ---
    for idx, q in enumerate(all_questions, start=1):
        question_text = q.get("question", "")
        doc.add_paragraph(f"Q{idx}: {question_text}", style='List Number')

        options = q.get("options")
        if options and isinstance(options, dict):
            # This is an MCQ
            for key in sorted(options.keys()):
                p = doc.add_paragraph(f"{key}. {options[key]}")
                # Bold the correct answer
                if key == q.get("answer"):
                    p.runs[0].font.bold = True
        else:
            # This is a True/False question
            answer_val = q.get("answer")
            answer = "True" if str(answer_val).lower() == "true" else "False"
            doc.add_paragraph(f"Answer: {answer}")

    # --- Add Validation Notes ---
    validation_notes = quiz_data.get("validation_notes")
    if validation_notes:
        doc.add_page_break()
        doc.add_heading("Validation Notes", level=1)
        doc.add_paragraph(validation_notes)

    # --- FIX: Add Fact-Checking Sources ---
    fact_checking_sources = quiz_data.get("fact_checking_sources")
    if fact_checking_sources: # Check if list is not None and not empty
        # Add page break if notes were not added, to keep sources separate
        if not validation_notes:
            doc.add_page_break()
            
        doc.add_heading("Fact-Checking Sources", level=1)
        for source in fact_checking_sources:
            # Add as a bullet point
            doc.add_paragraph(source, style='List Bullet')

    try:
        doc.save(file_path)
        return f"Quiz successfully saved as Word document at '{file_path}'."
    except Exception as e:
        return f"Error saving Word document: {e}"


word_writer_tool = FunctionTool(func=quiz_to_word_with_sources)